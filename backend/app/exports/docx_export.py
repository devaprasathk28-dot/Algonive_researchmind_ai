from docx import Document

def generate_docx_report(
    filename,
    title,
    summary,
    critique
):

    doc = Document()

    doc.add_heading(
        title,
        level=1
    )

    doc.add_heading(
        "AI Summary",
        level=2
    )

    doc.add_paragraph(summary)

    doc.add_heading(
        "Research Critique",
        level=2
    )

    doc.add_paragraph(critique)

    doc.save(filename)

    return filename
