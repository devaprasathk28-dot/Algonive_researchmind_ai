from docx import Document

def generate_docx_report(

    report_data,
    output_path

):

    document = Document()

    document.add_heading(

        "ResearchMind AI Report",

        level=1
    )

    document.add_paragraph(

        report_data["summary"]
    )

    document.save(
        str(output_path)
    )

    return output_path
